import os
from docx import Document
from docx.shared import Inches

def add_images_to_word(input_dir, output_file):
    # Create a Word document
    doc = Document()

    # Traverse the main directory and sort subdirectories and files
    all_images = []
    for root, _, files in sorted(os.walk(input_dir)):
        for file in sorted(files):  # Sort files for consistent order
            if file.endswith(".jpg"):
                all_images.append(os.path.join(root, file))

    # Initialize page tracking
    images_per_column = 10  # Adjust based on desired image size and spacing
    column_count = 0  # Track current column (0: left, 1: right)
    image_count = 0  # Track total images on the current page

    for image_path in all_images:
        if image_count == 0 or column_count == 0:  # Add a new table on a new page or column reset
            if image_count > 0:  # Add a page break if it's not the first page
                doc.add_page_break()
            table = doc.add_table(rows=0, cols=2)  # Create a new 2-column table
        
        # Add a new row if column_count is 0 (new left column entry)
        if column_count == 0:
            row_cells = table.add_row().cells

        # Add the image to the appropriate column
        try:
            row_cells[column_count].add_paragraph().add_run().add_picture(image_path, width=Inches(2.5))
        except Exception as e:
            print(f"Error adding image {image_path}: {e}")
        
        # Update counters
        column_count = (column_count + 1) % 2  # Toggle between 0 (left) and 1 (right)
        if column_count == 0:  # Increment image count only after filling both columns
            image_count += 1

        # Reset counters if the page is full
        if image_count >= images_per_column:
            column_count = 0
            image_count = 0

    # Save the document
    doc.save(output_file)
    print(f"Images have been added to {output_file}.")

# Usage
input_directory = "path/to/your/images"  # Replace with the path to the top-level directory
output_word_file = "output.docx"
add_images_to_word(input_directory, output_word_file)
